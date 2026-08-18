"""Convert non-PDF source files (images, txt, md, docx) to PDF.

The detection pipeline is PDF-native, so other formats are converted to
PDF once at ingestion and the converted file is fed through the existing
pipeline unchanged. Conversions use PDFium (via Pillow for raster decoding)
except DOCX, which needs the optional ``python-docx`` package for text
extraction.
"""

from __future__ import annotations

from pathlib import Path

import pypdfium2 as pdfium

from privacy_firewall.renderer.pdfium_draw import PageWriter

IMAGE_SUFFIXES: frozenset[str] = frozenset(
    {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp", ".gif"}
)
"""Raster image formats Pillow can decode and wrap into a PDF."""

TEXT_SUFFIXES: frozenset[str] = frozenset({".txt", ".md"})
"""Plain-text formats rendered onto PDF pages as-is."""

DOCX_SUFFIXES: frozenset[str] = frozenset({".docx"})
"""Word documents (text extracted via the optional ``python-docx``)."""

SUPPORTED_SUFFIXES: frozenset[str] = frozenset(
    {".pdf"} | IMAGE_SUFFIXES | TEXT_SUFFIXES | DOCX_SUFFIXES
)
"""Every file suffix the studio/ingestion layer accepts."""

_PAGE_WIDTH = 595.0  # A4 in points
_PAGE_HEIGHT = 842.0
_MARGIN = 50.0
_FONT_NAME = "cour"  # monospace: predictable wrapping and value alignment
_FONT_SIZE = 10.0
_LINE_HEIGHT = _FONT_SIZE * 1.4


class ConversionError(ValueError):
    """A source file could not be converted to PDF."""


def is_supported(path: Path | str) -> bool:
    """Whether *path*'s suffix is an accepted document format."""
    return Path(path).suffix.lower() in SUPPORTED_SUFFIXES


def needs_conversion(path: Path | str) -> bool:
    """Whether *path* is a supported format that must be converted first."""
    suffix = Path(path).suffix.lower()
    return suffix in SUPPORTED_SUFFIXES and suffix != ".pdf"


def convert_to_pdf(source: Path, dest: Path) -> Path:
    """Convert *source* to a PDF at *dest* (cached by modification time).

    Args:
        source: The input file (image, txt, md, or docx).
        dest: Where to write the converted PDF.

    Returns:
        *dest*, for chaining.

    Raises:
        ConversionError: If the format is unsupported, the file is
            unreadable/corrupt, or DOCX support is not installed.
    """
    source = Path(source)
    dest = Path(dest)
    suffix = source.suffix.lower()
    if not source.exists():
        msg = f"source file not found: {source}"
        raise ConversionError(msg)
    if dest.exists() and dest.stat().st_mtime >= source.stat().st_mtime:
        return dest  # up-to-date conversion already on disk

    if suffix in IMAGE_SUFFIXES:
        _image_to_pdf(source, dest)
    elif suffix in TEXT_SUFFIXES:
        text = source.read_text(encoding="utf-8", errors="replace")
        _text_to_pdf(text, dest)
    elif suffix in DOCX_SUFFIXES:
        _text_to_pdf(_extract_docx_text(source), dest)
    else:
        msg = f"unsupported file type: {suffix or '(no extension)'}"
        raise ConversionError(msg)
    return dest


def _image_to_pdf(source: Path, dest: Path) -> None:
    """Wrap a raster image into a single-page PDF (no text layer — OCR's job)."""
    try:
        from PIL import Image

        with Image.open(source) as img:
            frame = img.convert("RGB")
            width, height = frame.size
            bitmap = pdfium.PdfBitmap.from_pil(frame)
    except Exception as exc:
        msg = f"could not read image {source.name}: {exc}"
        raise ConversionError(msg) from exc

    doc = pdfium.PdfDocument.new()
    try:
        # One point per pixel keeps the page the image's natural size, which is
        # what the OCR pipeline expects when it scales bboxes back to points.
        page = doc.new_page(float(width), float(height))
        image = pdfium.PdfImage.new(doc)
        image.set_bitmap(bitmap)
        image.set_matrix(pdfium.PdfMatrix().scale(float(width), float(height)))
        page.insert_obj(image)
        page.gen_content()
        with dest.open("wb") as handle:
            doc.save(handle)
    except Exception as exc:
        msg = f"could not convert image {source.name}: {exc}"
        raise ConversionError(msg) from exc
    finally:
        doc.close()


def _text_to_pdf(text: str, dest: Path) -> None:
    """Render plain text onto paginated A4 pages with a real text layer."""
    doc = pdfium.PdfDocument.new()
    try:
        probe = doc.new_page(_PAGE_WIDTH, _PAGE_HEIGHT)
        writer = PageWriter(doc, probe, _PAGE_HEIGHT)
        max_width = _PAGE_WIDTH - 2 * _MARGIN
        char_width = writer.text_width("M", _FONT_NAME, _FONT_SIZE)
        chars_per_line = max(1, int(max_width / char_width)) if char_width else 80
        lines_per_page = max(1, int((_PAGE_HEIGHT - 2 * _MARGIN) / _LINE_HEIGHT))

        lines: list[str] = []
        for raw in text.splitlines() or [""]:
            raw = raw.replace("\t", "    ")
            lines.extend(_wrap_line(raw, chars_per_line))

        for index, start in enumerate(range(0, max(len(lines), 1), lines_per_page)):
            page = probe if index == 0 else doc.new_page(_PAGE_WIDTH, _PAGE_HEIGHT)
            writer = PageWriter(doc, page, _PAGE_HEIGHT)
            y = _MARGIN + _FONT_SIZE
            for line in lines[start : start + lines_per_page]:
                if line:
                    writer.insert_text(
                        (_MARGIN, y), line, fontsize=_FONT_SIZE, fontname=_FONT_NAME
                    )
                y += _LINE_HEIGHT
            writer.finalize()
        with dest.open("wb") as handle:
            doc.save(handle)
    finally:
        doc.close()


def _wrap_line(line: str, width: int) -> list[str]:
    """Wrap one logical line at *width* characters, breaking on spaces."""
    if len(line) <= width:
        return [line]
    wrapped: list[str] = []
    while len(line) > width:
        cut = line.rfind(" ", 1, width + 1)
        if cut <= 0:
            cut = width
        wrapped.append(line[:cut])
        line = line[cut:].lstrip(" ")
    wrapped.append(line)
    return wrapped


def _extract_docx_text(source: Path) -> str:
    """Pull paragraph and table text out of a DOCX file.

    Raises:
        ConversionError: If ``python-docx`` is missing or the file is
            not a valid DOCX document.
    """
    try:
        import docx
    except ImportError as exc:
        msg = "DOCX support requires the python-docx package: pip install python-docx"
        raise ConversionError(msg) from exc

    try:
        document = docx.Document(str(source))
    except Exception as exc:
        msg = f"could not read DOCX {source.name}: {exc}"
        raise ConversionError(msg) from exc

    parts: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)
